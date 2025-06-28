import base64
import json
from requests_toolbelt.multipart import decoder
import uuid
from Utils import get_postgresql_connection
from fastapi import FastAPI
from docx import Document
import csv
import io
import re
import boto3
import traceback

app = FastAPI()

s3 = boto3.client('s3')
BUCKET_NAME = 'awstraindata'

def lambda_handler(event, context):
    try:
        # Decode base64-encoded body (API Gateway encodes binary automatically)
        print(f"Received event")
        if event.get("isBase64Encoded", False):
            body = base64.b64decode(event['body'])
        else:
            body = event['body'].encode("utf-8")
        print(f"Decoded body length: {len(body)} bytes")
        # Get content-type header
        content_type = event['headers'].get('Content-Type') or event['headers'].get('content-type')
        if not content_type:
            return {"statusCode": 400, "body": "Missing Content-Type header"}

        # Parse multipart form
        multipart_data = decoder.MultipartDecoder(body, content_type)
        print(f"Multipart data parts: {len(multipart_data.parts)}")
        s3_urls = []
        conn = get_postgresql_connection()
        cursor = conn.cursor()
        for part in multipart_data.parts:
            print(f"Processing part: {part.headers.get(b'Content-Disposition')}")
            filename = part.headers.get(b'Content-Disposition').decode().split('filename="')[1].split('"')[0]
            file_stream = io.BytesIO(part.content)
            file_stream.seek(0)
            file_id = str(uuid.uuid4())
            s3_key = f"raw_data/{file_id}-{filename}"
            # Upload to S3
            s3.put_object(
                Bucket=BUCKET_NAME,
                Key=s3_key,
                Body=file_stream,
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            # Extract file name from content-disposition
            articles = extract_articles(io.BytesIO(part.content))
            print(f"Extracted {len(articles)} articles from part")
            for article in articles:
                print(f"Processing article: {article['Title']}")
                output_csv = io.StringIO()
                writer = csv.DictWriter(output_csv, fieldnames=["Title", "Source", "Date", "Content"])
                writer.writeheader()
                writer.writerow(article)
                article_id = str(uuid.uuid4())
                # Generate unique filename
                csv_filename = f"input/articles-{article_id}.csv"
                cursor.execute("""
                        INSERT INTO articles (article_id, title, body, source, published_date)
                        VALUES (%s, %s, %s, %s, %s)""", (article_id, article['Title'], article['Content'], article['Source'], article['Date']))
                # Upload to S3
                print(f"Uploading CSV to S3: {csv_filename}")
                s3.put_object(
                    Bucket=BUCKET_NAME,
                    Key=csv_filename,
                    Body=output_csv.getvalue(),
                    ContentType='text/csv'
                )
                conn.commit() 
                print(f"Uploaded CSV to S3: {csv_filename}")
                s3_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{csv_filename}"
                s3_urls.append(s3_url)
        cursor.close()
        conn.close()
        return {
            "statusCode": 200,
            "headers": {
                "Content-Type": "application/json"
            },
            "body": json.dumps({
                "status": "success",
                "s3_urls": s3_urls
            })
        }
    except Exception as e:
        traceback.print_exc() 
        return {"statusCode": 500, "body": f"❌ Error: {str(e)}"}

def extract_articles(file_stream):
    print(f"Extracting articles from file stream")
    doc = Document(file_stream)
    print(f"Document loaded with {len(doc.paragraphs)} paragraphs")
    text = "\n".join(p.text for p in doc.paragraphs)
    pattern = re.compile(
        r'Title:\s*(.*?)\s*Source:\s*(.*?)\s*Date:\s*(.*?)\s*(?=(?:\d{1,2}\)|Title:)|\Z)',
        re.DOTALL
    )
    matches = pattern.findall(text)
    print(f"Found {len(matches)} matches in the document")
    articles = []
    for match in matches:
        print(f"Processing match: {match}")
        title = match[0].strip()
        source = match[1].strip()
        date_parts = match[2].strip().split("\n", 1)
        date = date_parts[0].strip()
        content = date_parts[1].strip() if len(date_parts) > 1 else ""
        print(f"Extracted article - Title: {title}, Source: {source}, Date: {date}, Content length: {len(content)}")
        articles.append({
            "Title": title,
            "Source": source,
            "Date": date,
            "Content": content
            })
    return articles