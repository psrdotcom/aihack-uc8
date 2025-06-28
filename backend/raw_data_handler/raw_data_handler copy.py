import base64
import json
import time
import uuid
from Utils import get_postgresql_connection
from fastapi import FastAPI
from docx import Document
import csv
import io
import re
import boto3
import traceback

BUCKET_NAME = 'awstraindata'
role = 'arn:aws:iam::269854564686:role/hackathon-comprehend-role'
def lambda_handler(event, context):
    try:
        conn = get_postgresql_connection()
        cursor = conn.cursor()
        comprehend = boto3.client('comprehend', region_name='us-east-1')
        for record in event['Records']:
            print(f"New record: {record}")
            bucket = record['s3']['bucket']['name']
            key = record['s3']['object']['key']
            print(f"Processing file from bucket: {bucket}, key: {key}")
            s3 = boto3.client('s3')
            print(f"Connecting to S3 bucket: {bucket}")
            obj = s3.get_object(Bucket=bucket, Key=key)
            stream = io.BytesIO(obj['Body'].read()) 
            articles = extract_articles(stream)
            print(f"Extracted {len(articles)} articles from part")
            for article in articles:
                print(f"Processing article: {article['Title']}")
                output_csv = io.StringIO()
                writer = csv.DictWriter(output_csv, fieldnames=["Title", "Source", "Date", "Content"])
                # writer.writeheader()
                writer.writerow(article)
                article_id = str(uuid.uuid4())
                # Generate unique filename
                csv_filename = f"input/articles-{article_id}.csv"
                cursor.execute("""
                        INSERT INTO articles (article_id, title, body, source, published_date)
                        VALUES (%s, %s, %s, %s, %s)""", (article_id, article['Title'], article['Content'], article['Source'], article['Date']))
                # Upload to S3
                print(f"Uploading CSV to S3: {csv_filename}")
                s3.put_object(
                    Bucket=BUCKET_NAME,
                    Key=csv_filename,
                    Body=output_csv.getvalue(),
                    ContentType='text/csv'
                )
                conn.commit() 
                start_jobs(f's3://{BUCKET_NAME}/{csv_filename}', article_id, comprehend, role, cursor, conn)
        cursor.close()
        conn.close()
    except Exception as e:
        traceback.print_exc()
        print(f"Error processing event: {e}")
def get_data_inline(data, articles_id, comprehend, role_arn, cursor, conn):
    entities_response = comprehend.detect_entities(
                Text=data['Content'],
                DataAccessRoleArn=role_arn,
                LanguageCode='en'
            )
    add_entities_to_article(conn, cursor, articles_id, entities_response['Entities'])
    for entity in entities_response['Entities']:
        print(f"Found entity: {entity['Text']} (Type: {entity['Type']})")

def start_jobs(s3_uri, articles_id, comprehend, role_arn, cursor, conn):
    print(f"Starting jobs for article ID: {articles_id}")
    entities_job = comprehend.start_entities_detection_job(
                InputDataConfig={'S3Uri': s3_uri, 'InputFormat': 'ONE_DOC_PER_LINE'},
                OutputDataConfig={'S3Uri': 's3://awstraindata/output/entities/'},
                DataAccessRoleArn=role_arn,
                LanguageCode='en',
                JobName='MyEntityDetectionJob_'+ articles_id + '_' + str(int(time.time()))
            )
    print(f"Entities job started: {entities_job['JobId']}")
    result = comprehend.describe_entities_detection_job(JobId=entities_job['JobId'])
    print(f"Entities job description: {result}")
    entities_output = result['EntitiesDetectionJobProperties']['OutputDataConfig']['S3Uri']

    # SENTIMENT detection job
    sentiment_job = comprehend.start_sentiment_detection_job(
        InputDataConfig={'S3Uri': s3_uri, 'InputFormat': 'ONE_DOC_PER_LINE'},
        OutputDataConfig={'S3Uri': 's3://awstraindata/output/sentiment/'},
        DataAccessRoleArn=role_arn,
        LanguageCode='en',
        JobName='MySentimentDetectionJob_' + articles_id + '_' + str(int(time.time()))
    )
    print(f"Sentiment job started: {sentiment_job['JobId']}")
    res = comprehend.describe_sentiment_detection_job(JobId=sentiment_job['JobId'])
    print(f"Sentiment job description: {res}")
    sentiment_output = res['SentimentDetectionJobProperties']['OutputDataConfig']['S3Uri']

    # KEY PHRASES detection job
    phrases_job = comprehend.start_key_phrases_detection_job(
        InputDataConfig={'S3Uri': s3_uri, 'InputFormat': 'ONE_DOC_PER_LINE'},
        OutputDataConfig={'S3Uri': 's3://awstraindata/output/keyphrases/'},
        DataAccessRoleArn=role_arn,
        LanguageCode='en',
        JobName='MyKeyPhrasesDetectionJob_' + articles_id + '_' + str(int(time.time()))
    )
    print(f"Key Phrases job started: {phrases_job['JobId']}")
    res = comprehend.describe_key_phrases_detection_job(JobId=phrases_job['JobId'])
    print(f"Key Phrases job description: {res}")
    key_phrases_output = res['KeyPhrasesDetectionJobProperties']['OutputDataConfig']['S3Uri']
    print("Entities Job Response:", entities_output)
    print("Sentiment Job Response:", sentiment_output)
    print("Key Phrases Job Response:", key_phrases_output)
    print("Inserting into comprehend_jobs table")
    cursor.execute("""
        INSERT INTO comprehend_jobs (article_id, input_s3_uri, entities_path, sentiment_path, key_phrases_path)
        VALUES (%s, %s, %s, %s, %s)""", (articles_id, s3_uri, entities_output.replace('s3://awstraindata/', ''), sentiment_output.replace('s3://awstraindata/', ''), key_phrases_output.replace('s3://awstraindata/', '')))
    conn.commit()

def extract_articles(file_stream):
    print(f"Extracting articles from file stream")
    doc = Document(file_stream)
    print(f"Document loaded with {len(doc.paragraphs)} paragraphs")
    text = "\n".join(p.text for p in doc.paragraphs)
    pattern = re.compile(
        r'Title:\s*(.*?)\s*Source:\s*(.*?)\s*Date:\s*(.*?)\s*(?=(?:\d{1,2}\)|Title:)|\Z)',
        re.DOTALL
    )
    matches = pattern.findall(text)
    print(f"Found {len(matches)} matches in the document")
    articles = []
    for match in matches:
        print(f"Processing match: {match}")
        title = match[0].strip()
        source = match[1].strip()
        date_parts = match[2].strip().split("\n", 1)
        date = date_parts[0].strip()
        content = date_parts[1].strip() if len(date_parts) > 1 else ""
        print(f"Extracted article - Title: {title}, Source: {source}, Date: {date}, Content length: {len(content)}")
        articles.append({
            "Title": title,
            "Source": source,
            "Date": date,
            "Content": content
            })
    return articles


def add_entities_to_article(conn, cursor, article_id, entities):
    entities_text = [entity['Text'] for entity in entities]
    print(f"Entities to be added: {entities_text}")
    cursor.execute("SELECT * FROM entities WHERE entity in %s", (tuple(entities_text),))
    entity_db_array = cursor.fetchall()
    print(f"Entities in DB: {entity_db_array}")
    location_mentions = []
    officials_involved = []
    relevance_category = []
    print(f"article_id: {article_id}")
    
    print(f"Relevance category: {relevance_category}")
    for entity in entities:
        print(f"Processing entity: {entity}")
        entity_in_db = [db_entity for db_entity in entity_db_array if db_entity[3].lower() == entity['Text'].lower()]
        print(f"Entity in DB: {entity_in_db}")
        if not entity_in_db:
            current_time = datetime.datetime.utcnow()
            cursor.execute("INSERT INTO entities (create_time,entity,type) VALUES (%s, %s, %s) RETURNING id", (current_time, entity['Text'], entity['Type']))
            conn.commit()
            db_entity = cursor.fetchone()
            print(f"Inserted new entity: {db_entity}")
            if entity['Type'] == 'LOCATION':
                location_mentions.append(db_entity[0])
            elif entity['Type'] == 'PERSON':
                officials_involved.append(db_entity[0])
            else:
                relevance_category.append(db_entity[0])
        else:
            print(f"Entity already exists in DB: {entity_in_db}")
            if entity['Type'] == 'LOCATION':
                location_mentions.append(entity_in_db[0][0])
            elif entity['Type'] == 'PERSON':
                officials_involved.append(entity_in_db[0][0])
            else:
                relevance_category.append(entity_in_db[0][0])
    if location_mentions:
        location_mentions = ','.join(map(str, location_mentions))
        cursor.execute("""update articles set location_mentions = %s where article_id = %s""", (location_mentions, article_id))

    if officials_involved:
        officials_involved = ','.join(map(str, officials_involved))
        cursor.execute("""update articles set officials_involved = %s where article_id = %s""", (officials_involved, article_id))

    if relevance_category:
        cursor.execute("SELECT relevance_category FROM articles WHERE article_id = %s", (article_id,))
        existing = cursor.fetchone()
        relevance_category = ','.join(map(str, relevance_category))
        if existing[0] is not None:
            print(f"Existing relevance category: {existing[0]}")
            relevance_category = relevance_category + ',' + existing[0]
        cursor.execute("""update articles set relevance_category = %s where article_id = %s""", (relevance_category, article_id))
        cursor.execute("""update articles set relevance_category = %s where article_id = %s""", (relevance_category, article_id))