import base64
import datetime
import json
import time
import uuid
from Utils import get_postgresql_connection
from fastapi import FastAPI
from docx import Document
import csv
import io
import re
import boto3
import traceback
from Agent_helper import is_relevance
from botocore.config import Config

BUCKET_NAME = 'awstraindata'
role = 'arn:aws:iam::269854564686:role/hackathon-comprehend-role'
def lambda_handler(event, context):
    try:
        # conn = get_postgresql_connection()
        # cursor = conn.cursor()
        # comprehend = boto3.client('comprehend', region_name='us-east-1')
        # for record in event['Records']:
        #     print(f"New record: {record}")
        #     bucket = record['s3']['bucket']['name']
        #     key = record['s3']['object']['key']
        #     print(f"Processing file from bucket: {bucket}, key: {key}")
        #     s3 = boto3.client('s3')
        #     print(f"Connecting to S3 bucket: {bucket}")
        #     obj = s3.get_object(Bucket=bucket, Key=key)
        #     stream = io.BytesIO(obj['Body'].read()) 
        #     articles = extract_articles(stream)
        #     print(f"Extracted {len(articles)} articles from part")
        #     for article in articles:
        #         print(f"Processing article: {article['Title']}")
        #         # Check if article is relevant
        #         is_relevant = is_relevance(article)
        #         if not is_relevant:
        #             print(f"Article {article['Title']} is not relevant, skipping")
        #             continue
        #         output_csv = io.StringIO()
        #         writer = csv.DictWriter(output_csv, fieldnames=["Title", "Source", "Date", "Content"])
        #         # writer.writeheader()
        #         writer.writerow(article)
        #         article_id = str(uuid.uuid4())
        #         # Generate unique filename
        #         csv_filename = f"input/articles-{article_id}.csv"
        #         cursor.execute("""
        #                 INSERT INTO articles (article_id, title, body, source, published_date)
        #                 VALUES (%s, %s, %s, %s, %s)""", (article_id, article['Title'], article['Content'], article['Source'], article['Date']))
        #         # Upload to S3
        #         print(f"Uploading CSV to S3: {csv_filename}")
        #         conn.commit() 
        #         get_data_inline(output_csv.getvalue(), article_id, article['Date'], comprehend, cursor, conn)
        # cursor.close()
        # conn.close()
        config = Config(connect_timeout=10, read_timeout=30)
        lambda_client = boto3.client('lambda',config=config)
        print("Invoking second Lambda function")
        response = lambda_client.invoke(
            FunctionName='clustering_service',
            InvocationType='Event'
        )
        print(f"Second Lambda function invoked: {response}")
    except Exception as e:
        traceback.print_exc()
        print(f"Error processing event: {e}")

def get_data_inline(data, articles_id, article_date, comprehend, cursor, conn):
    print(f"Processing data for article ID: {articles_id}")
    entities_response = comprehend.detect_entities(
                Text=data,
                LanguageCode='en'
            )
    print(f"Entities detected: {entities_response['Entities']}")
    add_entities_to_article(conn, cursor, articles_id, entities_response['Entities'])
    response = comprehend.detect_key_phrases(
            Text=data,
            LanguageCode='en'
        )
    print(f"Key phrases detected: {response['KeyPhrases']}")
    for keyPhrase in response['KeyPhrases']:
        keyPhrase['Type'] = 'KeyPhrase'
    add_keyphrase_to_article(conn, cursor, articles_id, article_date, response['KeyPhrases'])
    sentiment_response = comprehend.detect_sentiment(
                Text=data,
                LanguageCode='en'
    )
    print(f"Sentiment detected: {sentiment_response['Sentiment']}")
    sentiment = sentiment_response['Sentiment']
    if sentiment:
        cursor.execute("""update articles set sentiment = %s where article_id = %s""", (sentiment, articles_id))

def extract_articles(file_stream):
    print(f"Extracting articles from file stream")
    doc = Document(file_stream)
    print(f"Document loaded with {len(doc.paragraphs)} paragraphs")
    text = "\n".join(p.text for p in doc.paragraphs)
    pattern = re.compile(
        r'Title:\s*(.*?)\s*Source:\s*(.*?)\s*Date:\s*(.*?)\s*(?=(?:\d{1,2}\)|Title:)|\Z)',
        re.DOTALL
    )
    matches = pattern.findall(text)
    print(f"Found {len(matches)} matches in the document")
    articles = []
    for match in matches:
        print(f"Processing match: {match}")
        title = match[0].strip()
        source = match[1].strip()
        date_parts = match[2].strip().split("\n", 1)
        date = date_parts[0].strip()
        content = date_parts[1].strip() if len(date_parts) > 1 else ""
        print(f"Extracted article - Title: {title}, Source: {source}, Date: {date}, Content length: {len(content)}")
        articles.append({
            "Title": title,
            "Source": source,
            "Date": date,
            "Content": content
            })
    return articles


def add_keyphrase_to_article(conn, cursor, article_id, article_date, entities):
    entities_text = [entity['Text'] for entity in entities]
    print(f"Entities to be added: {entities_text}")
    print(f"article_id: {article_id}")
    if entities_text:
        relevance_category = ','.join(map(str, entities_text))
        cursor.execute("""update articles set relevance_category = %s where article_id = %s""", (relevance_category, article_id))

def add_entities_to_article(conn, cursor, article_id, entities):
    entities_text = [entity['Text'] for entity in entities]
    print(f"Entities to be added: {entities_text}")
    location_mentions = []
    officials_involved = []
    print(f"article_id: {article_id}")
    for entity in entities:
        if entity['Type'] == 'LOCATION':
            location_mentions.append(entity['Text'].lower())
        elif entity['Type'] == 'PERSON' or entity['Type'] == 'ORGANIZATION':
            officials_involved.append(entity['Text'].lower())
    print(f"Processing entity: {entity}")
    if location_mentions:
        location_mentions = ','.join(map(str, location_mentions))
        cursor.execute("""update articles set location_mentions = %s where article_id = %s""", (location_mentions, article_id))

    if officials_involved:
        officials_involved = ','.join(map(str, officials_involved))
        cursor.execute("""update articles set officials_involved = %s where article_id = %s""", (officials_involved, article_id))